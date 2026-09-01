###   ===========================================================================
###                                      文件说明
###   ===========================================================================

# __file__ = "knowledge_vector.py"
# Description:此脚本实现了从指定文件夹中读取多种格式的文档（.docx、.doc、.pdf），将其内容进行切分，
# 并使用 DashScopeEmbeddings 进行文本嵌入，最后将嵌入向量存储到 Milvus 向量数据库中。
# 请参照.env文件(已提供)设置环境变量，包括 Milvus 的连接信息和 DashScope API Key。

import os
from langchain_core.documents import Document
from contextlib import redirect_stdout, redirect_stderr
from pathlib import Path
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from tqdm import tqdm
from vector.milvus_vector import MilvusVector
from langchain_community.embeddings import DashScopeEmbeddings


### 用于抑制函数输出的装饰器
def suppress_output(func):
    def wrapper(*args, **kwargs):
        with open(os.devnull, "w") as devnull:
            with redirect_stdout(devnull), redirect_stderr(devnull):
                return func(*args, **kwargs)

    return wrapper


### 读取文件内容的工具函数
@suppress_output
def read_file(file_path: str):
    if file_path.endswith(".docx"):
        return read_docx_file(file_path)
    elif file_path.endswith(".doc"):
        return read_doc_file(file_path)  # 使用相同的函数处理.doc文件
    elif file_path.endswith(".pdf"):
        return read_pdf_file(file_path)


### 使用python-docx库读取.doc或.docx文件内容
def read_docx_file(file_path: str) -> str:
    from docx import Document

    document = Document(file_path)
    content = "\n".join([para.text for para in document.paragraphs])

    return content


### doc需要单独处理为docx才可以继续处理
def read_doc_file(file_path: str) -> str:
    import subprocess
    import os

    # 创建一个临时的docx文件路径
    temp_docx_path = file_path + "x"

    # 使用libreoffice将.doc转换为.docx
    subprocess.run(
        [
            "libreoffice",
            "--headless",
            "--convert-to",
            "docx",
            file_path,
            "--outdir",
            os.path.dirname(file_path),
        ],
        check=True,
    )

    # 读取转换后的.docx文件内容
    content = read_docx_file(temp_docx_path)

    # 删除临时的.docx文件
    os.remove(temp_docx_path)

    return content


### 使用PyMuPDF库读取.pdf文件内容
def read_pdf_file(file_path: str) -> str:
    import fitz  # PyMuPDF

    doc = fitz.open(file_path)
    content = ""
    for page in doc:
        content += page.get_text()

    return content


def read_split_and_store(file_path, collection_name, dashscope_api_key):
    ### 初始化切分器
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,  # 每块目标长度（按字符或 token）
        chunk_overlap=50,  # 相邻块之间的重叠长度（缓解上下文断裂）
        length_function=len,  # 默认按字符计算长度；若用 token，需替换为 tokenizer
        separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", " ", ""],
    )

    ### 使用Path读取所有文件路径
    root_dir = Path(file_path)
    all_file_paths = list(root_dir.rglob("*.*"))
    if not all_file_paths:
        print(f"在路径 {file_path} 下未找到任何文件，请检查路径是否正确。")
        exit(1)
    docs = []

    for file_path in tqdm(all_file_paths, desc="Processing files"):
        file_name = file_path.name
        try:
            one_file_content = read_file(str(file_path))
        except Exception as e:
            print(f"读取文件 {file_name} 出错，跳过。错误信息：{e}")
            continue
        one_file_all_chunks = text_splitter.split_text(one_file_content)
        for chunk in one_file_all_chunks:
            metadata = {
                "source": str(file_name),
                "embed": "DashScopeEmbedding",
            }
            doc = Document(page_content=chunk, metadata=metadata)
            docs.append(doc)
    print("共计切块数量：", len(docs))

    print("文本嵌入并且向量存储到 Milvus...")

    embedding = DashScopeEmbeddings(
        model="text-embedding-v3", dashscope_api_key=dashscope_api_key
    )
    milvus_vector = MilvusVector(collection_name=collection_name, embedding=embedding)
    milvus_vector.vector_store(docs)

    print("文本嵌入并存储完成。存储collection名称：", collection_name)
